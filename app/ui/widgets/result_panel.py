"""ResultPanel — VLM解析结果展示 + 导出（Task 12 重构版：QTabBar 标签页切换视图）

设计要点：
- QTabBar（3 个标签：字段提取/Markdown预览/表格数据）替代原下拉框视图切换，
  标签页与 QStackedWidget 内容区通过 currentChanged 信号联动
- P0-a 结构化默认视图：字段提取为默认 Tab 0，Markdown 降为非默认原文视图
- 字段表格 3 列（字段/值/状态），状态文案 ✓ 已确认 / ⚠ 待确认 / ⚠ 冲突 / — 未找到，
  颜色走 ThemeManager 的 success/warning/error/text_disabled
- 数据源 page_result.structured.fields（StructuredResult）；兼容旧 2 参
  load_result(pr, finance_result) 的 FinanceResult 路径
- field_selected(row) 信号：字段行点击（Phase 4 hook）
- 导出按钮保留（📥 导出），支持 md/json/docx/xlsx 四种格式
- 全部颜色/字体/间距来自 ThemeManager，禁止硬编码颜色
  （原校验失败硬编码红色 QColor(0xd1,0x34,0x38) → 主题 error 色）
- 表格渲染兜底：tabulate 未安装时 df.to_markdown 抛 ImportError，降级为
  df.to_string 纯文本输出，避免运行期崩溃
"""
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QTextEdit, QTableWidget,
    QTableWidgetItem, QPushButton, QFileDialog, QStackedWidget,
    QTabBar, QMessageBox, QHeaderView, QAbstractItemView,
)
from PyQt6.QtCore import Qt, pyqtSignal as Signal
from PyQt6.QtGui import QColor
from typing import Optional
import json
import logging

from app.ui.theme_manager import ThemeManager
from app.ui.widgets.button_style import primary_qss
from app.models.page_result import StructuredResult, FinanceResult

logger = logging.getLogger("PDFOCR")

# 标签页顺序（索引即 content_stack 页面索引）：结构化字段为默认 Tab 0
TAB_LABELS = ["字段提取", "Markdown预览", "表格数据"]

# 结构化字段状态 → 展示文案 / 主题色角色（禁止硬编码颜色）
_STATUS_TEXT = {
    "confirmed": "✓ 已确认",
    "pending": "⚠ 待确认",
    "conflict": "⚠ 冲突",
    "not_found": "— 未找到",
}
_STATUS_COLOR = {
    "confirmed": "success",
    "pending": "warning_text",  # 文本用途 → 压暗版；圆点用途仍用 warning
    "conflict": "error",
    "not_found": "text_disabled",
}


class ResultPanel(QWidget):
    """右面板：字段提取 / Markdown预览 / 表格数据 + 导出（标签页切换）"""

    field_selected = Signal(int)  # 字段行点击（Phase 4 hook）

    def __init__(self, parent=None):
        super().__init__(parent)
        self._page_result = None
        self._structured = None  # StructuredResult | FinanceResult | None
        self._init_ui()
        # Task 15：主题切换后由 ThemeManager 触发重建 QSS
        ThemeManager.register_refresh_callback(self.apply_theme)

    def _init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # 标签页切换（替代原视图切换下拉框）
        self.tab_bar = QTabBar()
        for label in TAB_LABELS:
            self.tab_bar.addTab(label)
        layout.addWidget(self.tab_bar)

        # 内容区（堆叠视图）— 顺序必须与 TAB_LABELS 一致：
        # 索引0=字段提取 / 1=Markdown预览 / 2=表格数据
        self.content_stack = QStackedWidget()

        # 视图0: 字段提取表格（结构化默认视图）
        self._field_table = QTableWidget()
        self._field_table.setColumnCount(3)
        self._field_table.setHorizontalHeaderLabels(["字段", "值", "状态"])
        self._field_table.horizontalHeader().setStretchLastSection(True)
        self._field_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.ResizeToContents
        )
        self._field_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.Stretch
        )
        self._field_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self._field_table.setAlternatingRowColors(True)
        self._field_table.cellClicked.connect(self._on_field_clicked)
        self.content_stack.addWidget(self._field_table)

        # 视图1: Markdown预览（原文视图）
        self._md_view = QTextEdit()
        self._md_view.setReadOnly(True)
        self.content_stack.addWidget(self._md_view)

        # 视图2: 表格数据预览
        self._table_view = QTextEdit()
        self._table_view.setReadOnly(True)
        self.content_stack.addWidget(self._table_view)

        layout.addWidget(self.content_stack, stretch=1)

        # 导出按钮
        self.export_btn = QPushButton("📥 导出")
        self.export_btn.clicked.connect(self._on_export)
        layout.addWidget(self.export_btn, alignment=Qt.AlignmentFlag.AlignRight)

        # 标签页切换联动内容区
        self.tab_bar.currentChanged.connect(self._on_tab_changed)

        # 构造时烘焙样式（可安全重复执行）
        self.apply_theme()

    def apply_theme(self):
        """重建全部内嵌 QSS（Task 15：ThemeManager.set_theme 后调用）"""
        self.tab_bar.setStyleSheet(f"""
            QTabBar::tab {{
                background-color: {ThemeManager.get_color('bg_surface')};
                color: {ThemeManager.get_color('text_secondary')};
                padding: {ThemeManager.get_spacing('sm')}px
                         {ThemeManager.get_spacing('md')}px;
                border: none;
                border-bottom: 2px solid transparent;
            }}
            QTabBar::tab:selected {{
                color: {ThemeManager.get_color('primary')};
                border-bottom: 2px solid {ThemeManager.get_color('primary')};
            }}
            QTabBar::tab:hover:!selected {{
                background-color: {ThemeManager.get_color('bg_hover')};
            }}
            QTabBar::tab:focus {{
                /* 焦点环：键盘 Tab 聚焦的标签页显示焦点色下划线 */
                border-bottom: 2px solid {ThemeManager.get_color('border_focus')};
            }}
        """)
        # P1 字段表格：基础配色 + 焦点环（键盘导航可见性）
        self._field_table.setStyleSheet(f"""
            QTableWidget {{
                background-color: {ThemeManager.get_color('bg_surface')};
                border: none;
                outline: none;
                gridline-color: {ThemeManager.get_color('border')};
                alternate-background-color: {ThemeManager.get_color('bg_hover')};
            }}
            QTableWidget:focus {{
                border: 1px solid {ThemeManager.get_color('border_focus')};
            }}
            QTableWidget::item {{
                padding: {ThemeManager.get_spacing('xs')}px;
                color: {ThemeManager.get_color('text_primary')};
            }}
            QTableWidget::item:selected {{
                background-color: {ThemeManager.get_color('bg_selected')};
            }}
        """)
        # P2-a: 主操作按钮样式复用共享 single-source 样式（与 _btn_parse 一致）
        self.export_btn.setStyleSheet(primary_qss())

    # ---------- 视图切换 ----------

    def set_view(self, index_or_name):
        """切换当前视图（接受标签索引或标签名；TabBar 与内容区联动）"""
        if isinstance(index_or_name, str):
            if index_or_name not in TAB_LABELS:
                raise ValueError(f"未知视图: {index_or_name}")
            index_or_name = TAB_LABELS.index(index_or_name)
        self.tab_bar.setCurrentIndex(index_or_name)

    def _on_tab_changed(self, idx: int):
        """标签页切换：同步内容区并渲染当前视图"""
        self.content_stack.setCurrentIndex(idx)
        if self._page_result is None:
            return
        if idx == 0:
            # 字段提取（结构化默认视图）
            self._render_fields()
        elif idx == 1:
            # Markdown 原文视图
            self._md_view.setMarkdown(self._page_result.markdown or "(无内容)")
        elif idx == 2:
            # 表格数据
            if self._page_result.tables:
                texts = []
                for i, df in enumerate(self._page_result.tables):
                    try:
                        table_text = df.to_markdown(index=False)
                    except ImportError:
                        # tabulate 未安装时降级为纯文本表格，避免运行期崩溃
                        table_text = df.to_string(index=False)
                    texts.append(f"### 表格 {i+1}\n\n{table_text}")
                self._table_view.setMarkdown("\n\n".join(texts))
            else:
                self._table_view.setPlainText("(未检测到表格)")

    def _render_fields(self):
        """渲染字段表格：StructuredResult → 新状态语义；FinanceResult → 旧路径；None → 空"""
        st = self._structured
        if isinstance(st, StructuredResult):
            self._render_structured_fields(st)
        elif isinstance(st, FinanceResult):
            self._render_finance_fields(st)
        else:
            self._field_table.setRowCount(0)

    def _render_structured_fields(self, st: StructuredResult):
        """新状态语义：✓ 已确认 / ⚠ 待确认 / ⚠ 冲突 / — 未找到（主题色）"""
        fields = st.fields
        self._field_table.setRowCount(len(fields))
        for i, f in enumerate(fields):
            self._field_table.setItem(i, 0, QTableWidgetItem(f.label))
            self._field_table.setItem(i, 1, QTableWidgetItem(str(f.value or "")))
            item = QTableWidgetItem(_STATUS_TEXT.get(f.status, f.status))
            color_role = _STATUS_COLOR.get(f.status)
            if color_role:
                item.setForeground(QColor(ThemeManager.get_color(color_role)))
            self._field_table.setItem(i, 2, item)
        self._field_table.resizeColumnsToContents()

    def _render_finance_fields(self, st: FinanceResult):
        """兼容旧 2 参 load_result(pr, FinanceResult) 的 FinanceField 渲染路径"""
        if not st.fields:
            self._field_table.setRowCount(0)
            return
        self._field_table.setRowCount(len(st.fields))
        for i, f in enumerate(st.fields):
            self._field_table.setItem(i, 0, QTableWidgetItem(f.label))
            self._field_table.setItem(i, 1, QTableWidgetItem(str(f.value or "")))
            status = "✓" if f.validated else f"⚠ {f.validation_msg}"
            item = QTableWidgetItem(status)
            if not f.validated:
                item.setForeground(QColor(ThemeManager.get_color('error')))
            self._field_table.setItem(i, 2, item)
        self._field_table.resizeColumnsToContents()

    def _on_field_clicked(self, row: int, column: int):
        """字段行点击 → field_selected 信号（P1：结果 → 画布联动）"""
        self.field_selected.emit(row)

    def highlight_row(self, index: int):
        """选中并滚动到字段行（P1：画布点击 → 结果联动）"""
        if index < 0 or index >= self._field_table.rowCount():
            return
        self._field_table.selectRow(index)
        item = self._field_table.item(index, 0)
        if item is not None:
            self._field_table.scrollToItem(
                item, QAbstractItemView.ScrollHint.PositionAtCenter)

    # ---------- 数据接口 ----------

    def load_result(self, page_result, structured_result=None):
        """加载解析结果（兼容旧 2 参调用：结构化缺失时回退用 page_result.structured）"""
        self._page_result = page_result
        if structured_result is None and page_result.structured is not None:
            self._structured = page_result.structured
        else:
            # 可能为 None / StructuredResult / 旧 FinanceResult
            self._structured = structured_result
        self._on_tab_changed(self.tab_bar.currentIndex())

    def clear(self):
        """清空所有视图"""
        self._page_result = None
        self._structured = None
        self._md_view.clear()
        self._field_table.setRowCount(0)
        self._table_view.clear()

    # ---------- 导出 ----------

    def _on_export(self):
        """导出当前结果"""
        if self._page_result is None:
            QMessageBox.information(self, "提示", "暂无解析结果可导出")
            return
        filepath, _ = QFileDialog.getSaveFileName(
            self, "导出结果", "",
            "Markdown (*.md);;JSON (*.json);;Word (*.docx);;Excel (*.xlsx)"
        )
        if not filepath:
            return
        try:
            if filepath.endswith('.md'):
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(self._page_result.markdown or "")
            elif filepath.endswith('.json'):
                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(self._page_result.raw_json, f, ensure_ascii=False, indent=2)
            elif filepath.endswith('.docx'):
                from docx import Document
                doc = Document()
                doc.add_paragraph(self._page_result.markdown or "(空)")
                doc.save(filepath)
            elif filepath.endswith('.xlsx'):
                import pandas as pd
                if self._page_result.tables:
                    with pd.ExcelWriter(filepath) as writer:
                        for i, df in enumerate(self._page_result.tables):
                            df.to_excel(writer, sheet_name=f"Table_{i+1}", index=False)
                else:
                    pd.DataFrame().to_excel(filepath, index=False)
            QMessageBox.information(self, "导出成功", f"已保存到:\n{filepath}")
        except Exception as e:
            logger.error(f"导出失败: {e}", exc_info=True)
            QMessageBox.critical(self, "导出失败", str(e))
